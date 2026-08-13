"""Text extraction from uploaded documents.

Deferred from Phase 0 and landed here because Project 1 is the first thing that
needs it (CLAUDE.md Section 9: "optional documents" as SOP input).

Both libraries are optional dependencies, imported lazily. A missing library
produces an error naming what to install, rather than an ImportError at startup
or a silently empty result.
"""

from __future__ import annotations

import io
from enum import StrEnum
from pathlib import Path

from aiops_utils import ConfigurationError, ValidationError


class DocumentKind(StrEnum):
    PDF = "pdf"
    DOCX = "docx"
    TEXT = "text"
    MARKDOWN = "markdown"


_EXTENSIONS: dict[str, DocumentKind] = {
    ".pdf": DocumentKind.PDF,
    ".docx": DocumentKind.DOCX,
    ".txt": DocumentKind.TEXT,
    ".md": DocumentKind.MARKDOWN,
}

#: Guards against a caller passing a huge file straight into an AI prompt.
MAX_EXTRACTED_CHARS = 100_000


def detect_kind(filename: str) -> DocumentKind:
    """Determine the document type from its filename."""
    suffix = Path(filename).suffix.lower()
    kind = _EXTENSIONS.get(suffix)
    if kind is None:
        supported = ", ".join(sorted(_EXTENSIONS))
        raise ValidationError(
            f"{suffix or 'This file type'} is not supported. Upload one of: {supported}.",
            user_message=(
                f"That file type is not supported. Please upload a "
                f"{', '.join(sorted(e.lstrip('.').upper() for e in _EXTENSIONS))} file."
            ),
        )
    return kind


def _extract_pdf(data: bytes) -> str:
    try:
        import pymupdf
    except ImportError as exc:
        raise ConfigurationError(
            "Reading PDFs needs PyMuPDF. Install it with: pip install pymupdf",
            user_message="PDF reading is not available on this server.",
        ) from exc

    try:
        with pymupdf.open(stream=data, filetype="pdf") as document:
            return "\n\n".join(page.get_text() for page in document)
    except Exception as exc:  # noqa: BLE001 — a corrupt file is user error
        raise ValidationError(
            f"Could not read this PDF: {exc}",
            user_message="That PDF could not be read. It may be corrupt or password-protected.",
        ) from exc


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:
        raise ConfigurationError(
            "Reading Word files needs python-docx. Install it with: pip install python-docx",
            user_message="Word document reading is not available on this server.",
        ) from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 — a corrupt file is user error
        raise ValidationError(
            f"Could not read this Word document: {exc}",
            user_message="That Word document could not be read. It may be corrupt.",
        ) from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # Tables often hold the actual procedure in operational documents, so they
    # are extracted too rather than dropped.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(data: bytes, filename: str) -> str:
    """Extract plain text from an uploaded document.

    Raises `ValidationError` for an unsupported, empty, or unreadable file.
    """
    if not data:
        raise ValidationError(
            "The uploaded file is empty.",
            user_message="That file appears to be empty.",
        )

    kind = detect_kind(filename)
    if kind is DocumentKind.PDF:
        text = _extract_pdf(data)
    elif kind is DocumentKind.DOCX:
        text = _extract_docx(data)
    else:
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            # Fall back rather than fail — operational files are often written
            # on Windows in a legacy encoding.
            text = data.decode("latin-1", errors="replace")

    text = text.strip()
    if not text:
        raise ValidationError(
            f"No text could be extracted from {filename!r}.",
            user_message=(
                "No text could be read from that file. If it is a scanned PDF, "
                "the text needs to be selectable — scanned images are not supported."
            ),
        )

    if len(text) > MAX_EXTRACTED_CHARS:
        text = text[:MAX_EXTRACTED_CHARS]

    return text
